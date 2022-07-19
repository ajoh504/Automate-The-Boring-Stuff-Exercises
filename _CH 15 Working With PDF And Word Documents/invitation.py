#!python3
# invitation.py - Open a blank Word document that contains pre-set styles,
#                 and create a list of invitations for different guests. The
#                 guest names are stored in a .txt file.
#
#          USAGE: sys.argv[1] = .txt file containing guest names
#                 sys.argv[2] = blank .docx file with pre-set styles
#                 using Python version 3.8.10 and python-docx version 0.8.10
# changed: 1. converted program to class 2. moved to single function 3. added docstring for make_invitations() 4. used split() method and list comprehension in guests_list() function to remove any line break characters

import sys
import docx


class CustomInvitations:
    def __init__(self):
        """
        self.guests_lists: Callable function to return a list of lines from
        sys.argv[1] and remove any "\n" characters.
        """
        self.guests_list = lambda: [
            line.split("\n")[0]
            for line in open(self.guest_names, "r").readlines()
            if "\n" in line
        ]
        self.guest_names = sys.argv[1]
        self.blank_docx = sys.argv[2]
        self.invitation = (
            "It would be a pleasure to have the company of",
            "",  # placeholder for guest name
            "at 11010 Memory Lane on the Evening of",
            "April 1st",
            "at 7 o'clock",
        )



    def make_invitations(self):
        """
        Loop through guests_lists() and invitation. Place guest's name at index 1.
        For all other cases, add "line" with the appropriate Custom Style. Keep
        track of paragraph count in order to determine where the line break is
        placed.
        """
        doc = docx.Document(self.blank_docx)
        paragraph_count = 0
        for i, guest in enumerate(self.guests_list()):
            for j, line in enumerate(self.invitation):
                if j == 1:
                    doc.add_paragraph(guest, "Custom Style 2")
                elif j % 2 == 0:
                    doc.add_paragraph(line, "Custom Style 1")
                elif j % 2 != 0:
                    doc.add_paragraph(line, "Custom Style 2")
                paragraph_count += 1
            doc.paragraphs[paragraph_count].runs[0].add_break(
                docx.enum.text.WD_BREAK.PAGE
            )
        doc.save("invitations.docx")


if __name__ == "__main__":
    custom_invitations = CustomInvitations()
    custom_invitations.make_invitations()
